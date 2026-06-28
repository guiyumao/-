from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSET_DIR = DOCS_DIR / "generated_assets"
OUTPUT_IMAGE = ASSET_DIR / "detailed_role_usecase_flow.png"
OUTPUT_DOCX = DOCS_DIR / "detailed_role_usecase_flow.docx"

BLUE = "#3f6ecb"
BLUE_DARK = "#284f9f"
LINE = "#2f5fb3"
TEXT = "#1d3f8f"
LANE = "#f7faff"
LANE_BORDER = "#c7d7f6"
WARN_BG = "#fff1c4"
WARN_TEXT = "#b56b00"
MUTED = "#5b6475"


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/Deng.ttf"),
    ]
    for font_path in candidates:
        if font_path.exists():
            return ImageFont.truetype(str(font_path), size)
    return ImageFont.load_default()


def center_text(
    draw: ImageDraw.ImageDraw,
    xy: tuple[float, float],
    text: str,
    font: ImageFont.ImageFont,
    fill: str = "black",
) -> None:
    box = draw.textbbox((0, 0), text, font=font)
    width = box[2] - box[0]
    height = box[3] - box[1]
    draw.text((xy[0] - width / 2, xy[1] - height / 2), text, font=font, fill=fill)


def draw_actor(draw: ImageDraw.ImageDraw, x: int, y: int, label: str) -> None:
    label_font = load_font(22, bold=True)
    draw.ellipse([x - 13, y - 50, x + 13, y - 24], fill=LINE)
    draw.rounded_rectangle([x - 21, y - 22, x + 21, y + 38], radius=6, fill=LINE)
    draw.rectangle([x - 42, y - 15, x - 21, y + 24], fill=LINE)
    draw.rectangle([x + 21, y - 15, x + 42, y + 24], fill=LINE)
    draw.rectangle([x - 15, y + 38, x - 4, y + 70], fill=LINE)
    draw.rectangle([x + 4, y + 38, x + 15, y + 70], fill=LINE)
    center_text(draw, (x, y + 88), label, label_font, TEXT)


def draw_usecase(draw: ImageDraw.ImageDraw, x: int, y: int, text: str, w: int = 190, h: int = 58) -> None:
    font = load_font(21, bold=True)
    draw.ellipse([x - w / 2, y - h / 2, x + w / 2, y + h / 2], fill=BLUE, outline=BLUE_DARK, width=2)
    center_text(draw, (x, y), text, font, "white")


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    label: str | None = None,
    color: str = LINE,
    width: int = 2,
) -> None:
    sx, sy = start
    ex, ey = end
    draw.line([start, end], fill=color, width=width)
    dx = ex - sx
    dy = ey - sy
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux = dx / length
    uy = dy / length
    size = 10
    left = (ex - ux * size - uy * size * 0.6, ey - uy * size + ux * size * 0.6)
    right = (ex - ux * size + uy * size * 0.6, ey - uy * size - ux * size * 0.6)
    draw.polygon([end, left, right], fill=color)
    if label:
        font = load_font(16)
        center_text(draw, ((sx + ex) / 2, (sy + ey) / 2 - 16), label, font, color)


def draw_warning(draw: ImageDraw.ImageDraw, width: int) -> None:
    font = load_font(20)
    draw.rectangle([80, 18, width - 80, 62], fill=WARN_BG)
    draw.text((100, 30), "△ 多角色用例流程说明图需覆盖申请、审批、管理、运维、统计和权限配置等核心业务。", font=font, fill=WARN_TEXT)


def draw_lane(
    draw: ImageDraw.ImageDraw,
    y: int,
    role: str,
    description: str,
    usecases: list[str],
    actor_x: int = 120,
    first_x: int = 405,
    gap: int = 205,
) -> None:
    desc_font = load_font(18)
    lane_h = 172
    draw.rounded_rectangle([45, y - 78, 2355, y + lane_h - 78], radius=18, fill=LANE, outline=LANE_BORDER, width=3)
    draw_actor(draw, actor_x, y, role)
    draw.text((220, y - 52), description, font=desc_font, fill=MUTED)

    previous: tuple[int, int] | None = None
    for index, item in enumerate(usecases):
        x = first_x + index * gap
        w = 210 if len(item) <= 6 else 245
        draw_usecase(draw, x, y, item, w=w)
        if index == 0:
            arrow(draw, (actor_x + 54, y), (x - w // 2, y))
        elif previous:
            arrow(draw, (previous[0] + 103, y), (x - w // 2, y), label="流程")
        previous = (x, y)


def draw_cross_relation(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], label: str) -> None:
    arrow(draw, start, end, label=label, color="#6a7fc2", width=2)


def draw_diagram(path: Path) -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 2400, 2050
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    title_font = load_font(42, bold=True)
    note_font = load_font(19)

    draw_warning(draw, width)
    center_text(draw, (width / 2, 102), "实验室设备与耗材管理系统多角色详细用例流程说明图", title_font, TEXT)

    lanes = [
        (
            240,
            "学生",
            "授权范围内提交申请，主要查看个人记录和申请进度。",
            ["登录认证", "个人中心", "查看个人记录", "设备借用申请", "耗材申领申请", "危化品申请"],
        ),
        (
            430,
            "教师",
            "教学/科研业务发起人，可指导学生并提交设备、耗材、危化品相关申请。",
            ["登录认证", "指导学生申请", "设备借用申请", "耗材申领申请", "危化品领用申请", "设备报修"],
        ),
        (
            620,
            "实验室主任",
            "实验室业务负责人，负责审批、复核、调度和统计。",
            ["审批申请", "安全资质核验", "双人复核", "归还验收", "统计报表", "到期提醒"],
        ),
        (
            810,
            "设备管理员",
            "负责设备基础数据、借还确认、维修调度和逾期催还。",
            ["设备分类管理", "设备台账维护", "借用登记确认", "归还登记", "维修派单", "逾期催还"],
        ),
        (
            1000,
            "耗材管理员",
            "负责耗材分类、台账、入库、出库、库存批次和预警。",
            ["耗材分类管理", "耗材台账维护", "耗材入库", "出库确认", "批次库存维护", "安全库存预警"],
        ),
        (
            1190,
            "危化品管理员",
            "负责危化品台账、安全信息、领用、归还、废液和安全追踪。",
            ["危化品台账", "CAS/MSDS维护", "领用登记", "余量归还", "废液处理", "安全审计"],
        ),
        (
            1380,
            "维修人员",
            "负责故障处理过程、维修费用、维修结果登记。",
            ["接收维修任务", "记录维修过程", "登记维修费用", "提交维修结果", "等待主任验收"],
        ),
        (
            1570,
            "校准人员",
            "负责设备校准任务、证书、有效期和校准结果维护。",
            ["接收校准任务", "登记证书编号", "填写校准日期", "记录校准结果", "更新下次校准"],
        ),
        (
            1760,
            "系统管理员",
            "负责账号、角色、菜单权限、日志和系统配置。",
            ["用户管理", "角色权限管理", "菜单权限管理", "系统配置", "审计日志", "数据备份"],
        ),
    ]

    for lane in lanes:
        draw_lane(draw, *lane)

    # Cross-role include/extend/supervision relationships.
    draw_cross_relation(draw, (1520, 240), (405, 620), "危化品申请需审批")
    draw_cross_relation(draw, (1520, 430), (405, 620), "教师申请需审批")
    draw_cross_relation(draw, (1430, 620), (1425, 810), "异常设备转维修")
    draw_cross_relation(draw, (1220, 810), (405, 1380), "派单")
    draw_cross_relation(draw, (1425, 1570), (1430, 620), "到期提醒")
    draw_cross_relation(draw, (1430, 1000), (1430, 620), "库存预警上报")
    draw_cross_relation(draw, (1425, 1190), (815, 620), "双人复核/安全审核")
    draw_cross_relation(draw, (1015, 1760), (1015, 620), "权限支撑")

    draw.text(
        (70, 1985),
        "说明：本图按角色拆分职责。学生、教师侧重业务申请；主任侧重审批与验收；设备/耗材/危化品管理员负责专项台账和库存；维修、校准人员负责专业处理；系统管理员负责权限与审计。",
        font=note_font,
        fill="#555555",
    )

    image.save(path)


def build_docx() -> None:
    draw_diagram(OUTPUT_IMAGE)
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("实验室设备与耗材管理系统多角色详细用例流程说明图")
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(18)

    doc.add_picture(str(OUTPUT_IMAGE), width=Cm(26.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = caption.add_run("图 1 多角色详细用例流程说明图")
    cap.font.name = "Arial"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cap.font.size = Pt(11)

    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
    print(OUTPUT_IMAGE)
    print(OUTPUT_DOCX)
