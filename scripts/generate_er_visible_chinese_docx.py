from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSET_DIR = DOCS_DIR / "generated_assets"
BASE_DOCX = DOCS_DIR / "database_and_requirements_analysis_clean.docx"
OUTPUT_DOCX = DOCS_DIR / "database_and_requirements_analysis_er_chinese_visible.docx"
OUTPUT_IMAGE = ASSET_DIR / "system_er_diagram_chinese.png"


def cn(text: str) -> str:
    return text.encode("utf-8").decode("unicode_escape")


def load_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    font_candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/Deng.ttf"),
    ]
    for font_path in font_candidates:
        if font_path.exists():
            return ImageFont.truetype(str(font_path), size)
    return ImageFont.load_default()


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    center: tuple[float, float],
    text: str,
    font: ImageFont.FreeTypeFont | ImageFont.ImageFont,
    fill: str = "black",
) -> None:
    bbox = draw.textbbox((0, 0), text, font=font)
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.text((center[0] - width / 2, center[1] - height / 2), text, font=font, fill=fill)


def draw_er_image(path: Path) -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    width, height = 2600, 1650
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    title_font = load_font(46)
    section_font = load_font(28)
    entity_font = load_font(30)
    relation_font = load_font(24)
    cardinality_font = load_font(22)

    entities = {
        cn(r"\u5b9e\u9a8c\u5ba4"): (240, 420),
        cn(r"\u7528\u6237"): (240, 220),
        cn(r"\u89d2\u8272"): (620, 220),
        cn(r"\u83dc\u5355"): (1000, 220),
        cn(r"\u7cfb\u7edf\u65e5\u5fd7"): (620, 420),
        cn(r"\u8bbe\u5907\u7c7b\u522b"): (240, 700),
        cn(r"\u8bbe\u5907"): (620, 700),
        cn(r"\u8bbe\u5907\u501f\u7528"): (240, 980),
        cn(r"\u8bbe\u5907\u7ef4\u4fee"): (620, 980),
        cn(r"\u8bbe\u5907\u6821\u51c6"): (1000, 980),
        cn(r"\u8017\u6750\u7c7b\u522b"): (1380, 700),
        cn(r"\u8017\u6750"): (1760, 700),
        cn(r"\u8017\u6750\u5165\u5e93"): (1380, 980),
        cn(r"\u8017\u6750\u51fa\u5e93"): (1760, 980),
        cn(r"\u5371\u5316\u54c1"): (2260, 700),
        cn(r"\u5371\u5316\u54c1\u4f7f\u7528"): (2260, 980),
        cn(r"\u5e93\u5b58"): (1760, 1260),
    }

    entity_width = 220
    entity_height = 78
    relation_width = 130
    relation_height = 70

    def draw_section(label: str, box: tuple[int, int, int, int]) -> None:
        x1, y1, x2, y2 = box
        draw.rounded_rectangle([x1, y1, x2, y2], radius=18, outline="#c9d6ee", width=4)
        draw.text((x1 + 24, y1 + 18), label, font=section_font, fill="#335f9f")

    def draw_entity(label: str, center: tuple[int, int]) -> None:
        x, y = center
        box = [
            x - entity_width // 2,
            y - entity_height // 2,
            x + entity_width // 2,
            y + entity_height // 2,
        ]
        draw.rectangle(box, outline="#1f3fa3", width=3)
        draw_centered_text(draw, center, label, entity_font)

    def draw_relation(label: str, center: tuple[float, float]) -> None:
        x, y = center
        points = [
            (x, y - relation_height / 2),
            (x + relation_width / 2, y),
            (x, y + relation_height / 2),
            (x - relation_width / 2, y),
        ]
        draw.polygon(points, outline="#1f7a36", width=3)
        draw_centered_text(draw, center, label, relation_font)

    def edge_point(
        source: tuple[float, float],
        target: tuple[float, float],
        half_width: float,
        half_height: float,
    ) -> tuple[float, float]:
        sx, sy = source
        tx, ty = target
        dx = tx - sx
        dy = ty - sy
        if dx == 0 and dy == 0:
            return sx, sy
        if abs(dx) * half_height > abs(dy) * half_width:
            x = half_width if dx > 0 else -half_width
            y = dy * abs(x / dx)
        else:
            y = half_height if dy > 0 else -half_height
            x = dx * abs(y / dy)
        return sx + x, sy + y

    def connect(
        left: str,
        relation: str,
        right: str,
        left_cardinality: str,
        right_cardinality: str,
        offset: tuple[int, int] = (0, 0),
    ) -> None:
        x1, y1 = entities[left]
        x2, y2 = entities[right]
        mx = (x1 + x2) / 2 + offset[0]
        my = (y1 + y2) / 2 + offset[1]
        relation_center = (mx, my)

        draw_relation(relation, relation_center)

        start = edge_point((x1, y1), relation_center, entity_width / 2, entity_height / 2)
        rel_left = edge_point(relation_center, (x1, y1), relation_width / 2, relation_height / 2)
        rel_right = edge_point(relation_center, (x2, y2), relation_width / 2, relation_height / 2)
        end = edge_point((x2, y2), relation_center, entity_width / 2, entity_height / 2)

        draw.line([start, rel_left], fill="#333333", width=2)
        draw.line([rel_right, end], fill="#333333", width=2)
        draw_centered_text(
            draw,
            ((start[0] + rel_left[0]) / 2, (start[1] + rel_left[1]) / 2 - 18),
            left_cardinality,
            cardinality_font,
        )
        draw_centered_text(
            draw,
            ((rel_right[0] + end[0]) / 2, (rel_right[1] + end[1]) / 2 - 18),
            right_cardinality,
            cardinality_font,
        )

    def connect_labeled(
        left: str,
        right: str,
        label: str,
        left_cardinality: str,
        right_cardinality: str,
        via: tuple[int, int] | None = None,
    ) -> None:
        """Draw a lightweight labeled ER relationship to keep the total diagram readable."""
        x1, y1 = entities[left]
        x2, y2 = entities[right]
        start = edge_point((x1, y1), (x2, y2), entity_width / 2, entity_height / 2)
        end = edge_point((x2, y2), (x1, y1), entity_width / 2, entity_height / 2)
        if via:
            points = [start, via, end]
            lx, ly = via
        else:
            points = [start, end]
            lx = (start[0] + end[0]) / 2
            ly = (start[1] + end[1]) / 2
        draw.line(points, fill="#3f3f3f", width=2)
        draw_centered_text(draw, (lx, ly - 24), label, relation_font, "#1f7a36")
        draw_centered_text(draw, (start[0], start[1] - 20), left_cardinality, cardinality_font)
        draw_centered_text(draw, (end[0], end[1] - 20), right_cardinality, cardinality_font)

    draw_section(cn(r"\u6743\u9650\u4e0e\u7528\u6237"), (70, 130, 1130, 510))
    draw_section(cn(r"\u8bbe\u5907\u7ba1\u7406"), (70, 610, 1130, 1080))
    draw_section(cn(r"\u8017\u6750\u4e0e\u5371\u5316\u54c1"), (1210, 610, 2490, 1080))
    draw_section(cn(r"\u7edf\u4e00\u5e93\u5b58"), (1210, 1160, 2490, 1370))

    for entity, center in entities.items():
        draw_entity(entity, center)

    one = "1"
    many = "N"
    m = "M"
    connect(cn(r"\u7528\u6237"), cn(r"\u62e5\u6709"), cn(r"\u89d2\u8272"), m, many, (0, -72))
    connect(cn(r"\u89d2\u8272"), cn(r"\u6388\u6743"), cn(r"\u83dc\u5355"), m, many, (0, -72))
    connect(cn(r"\u5b9e\u9a8c\u5ba4"), cn(r"\u5305\u542b"), cn(r"\u7528\u6237"), one, many, (-78, 0))
    connect(cn(r"\u7528\u6237"), cn(r"\u4ea7\u751f"), cn(r"\u7cfb\u7edf\u65e5\u5fd7"), one, many, (0, 48))
    connect(cn(r"\u8bbe\u5907\u7c7b\u522b"), cn(r"\u5206\u7c7b"), cn(r"\u8bbe\u5907"), one, many, (0, -60))
    connect(cn(r"\u8bbe\u5907"), cn(r"\u501f\u7528"), cn(r"\u8bbe\u5907\u501f\u7528"), one, many, (-210, 0))
    connect(cn(r"\u8bbe\u5907"), cn(r"\u7ef4\u4fee"), cn(r"\u8bbe\u5907\u7ef4\u4fee"), one, many, (0, 40))
    connect(cn(r"\u8bbe\u5907"), cn(r"\u6821\u51c6"), cn(r"\u8bbe\u5907\u6821\u51c6"), one, many, (210, 0))
    connect(cn(r"\u8017\u6750\u7c7b\u522b"), cn(r"\u5206\u7c7b"), cn(r"\u8017\u6750"), one, many, (0, -60))
    connect(cn(r"\u8017\u6750"), cn(r"\u5165\u5e93"), cn(r"\u8017\u6750\u5165\u5e93"), one, many, (-200, 0))
    connect(cn(r"\u8017\u6750"), cn(r"\u51fa\u5e93"), cn(r"\u8017\u6750\u51fa\u5e93"), one, many, (0, 38))
    connect(cn(r"\u5371\u5316\u54c1"), cn(r"\u4f7f\u7528"), cn(r"\u5371\u5316\u54c1\u4f7f\u7528"), one, many, (0, 38))

    connect_labeled(cn(r"\u5b9e\u9a8c\u5ba4"), cn(r"\u8bbe\u5907"), cn(r"\u62e5\u6709"), one, many, (420, 610))
    connect_labeled(cn(r"\u5b9e\u9a8c\u5ba4"), cn(r"\u8017\u6750"), cn(r"\u62e5\u6709"), one, many, (1030, 620))
    connect_labeled(cn(r"\u5b9e\u9a8c\u5ba4"), cn(r"\u5371\u5316\u54c1"), cn(r"\u62e5\u6709"), one, many, (1180, 560))
    connect_labeled(cn(r"\u5b9e\u9a8c\u5ba4"), cn(r"\u5e93\u5b58"), cn(r"\u7ba1\u7406"), one, many, (980, 1240))
    connect_labeled(cn(r"\u7528\u6237"), cn(r"\u8bbe\u5907\u501f\u7528"), cn(r"\u7533\u8bf7"), one, many, (120, 630))
    connect_labeled(cn(r"\u8017\u6750"), cn(r"\u5e93\u5b58"), cn(r"\u5f62\u6210\u6279\u6b21"), one, many, (1680, 1100))
    connect_labeled(cn(r"\u5371\u5316\u54c1"), cn(r"\u5e93\u5b58"), cn(r"\u5f62\u6210\u6279\u6b21"), one, many, (2050, 1140))

    draw_centered_text(
        draw,
        (width / 2, 65),
        cn(r"\u5b9e\u9a8c\u5ba4\u8bbe\u5907\u4e0e\u8017\u6750\u7ba1\u7406\u7cfb\u7edf\u603b\u4f53 E-R \u56fe\uff08Chen \u6a21\u578b\uff09"),
        title_font,
    )
    image.save(path)


def remove_existing_images(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        if paragraph._p.xpath(".//w:drawing"):
            paragraph._element.getparent().remove(paragraph._element)
            continue
        text = paragraph.text.strip()
        if "\uff1f\uff1f" in text or "????" in text:
            paragraph._element.getparent().remove(paragraph._element)


def insert_er_image(doc: Document, image_path: Path) -> None:
    anchor_index = None
    for index, paragraph in enumerate(doc.paragraphs):
        if "Chen E-R" in paragraph.text and "\u6982\u5ff5\u7ed3\u6784" in paragraph.text:
            anchor_index = index
            break
    if anchor_index is None:
        raise RuntimeError("Could not find ER intro paragraph.")

    anchor = doc.paragraphs[anchor_index]

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(cn(r"\u7cfb\u7edf\u603b\u4f53 E-R \u56fe"))
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), cn(r"\u9ed1\u4f53"))
    title_run.font.size = Pt(14)

    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(image_path), width=Cm(15.2))

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption.add_run(cn(r"\u56fe 1 \u7cfb\u7edf\u603b\u4f53 E-R \u56fe\uff08Chen \u6a21\u578b\uff09"))
    caption_run.font.name = "Arial"
    caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), cn(r"\u5b8b\u4f53"))
    caption_run.font.size = Pt(10.5)

    for paragraph in [title, image_paragraph, caption]:
        paragraph._p.getparent().remove(paragraph._p)

    anchor._p.addnext(caption._p)
    anchor._p.addnext(image_paragraph._p)
    anchor._p.addnext(title._p)


def main() -> None:
    draw_er_image(OUTPUT_IMAGE)
    doc = Document(BASE_DOCX)
    remove_existing_images(doc)
    insert_er_image(doc, OUTPUT_IMAGE)
    doc.save(OUTPUT_DOCX)

    check = Document(OUTPUT_DOCX)
    print(OUTPUT_IMAGE)
    print(OUTPUT_DOCX)
    print(f"inline_shapes={len(check.inline_shapes)}")


if __name__ == "__main__":
    main()
