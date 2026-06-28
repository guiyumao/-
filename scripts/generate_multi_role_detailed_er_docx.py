from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSET_DIR = DOCS_DIR / "generated_assets"
OUTPUT_DOCX = DOCS_DIR / "multi_role_detailed_er.docx"
PERMISSION_IMAGE = ASSET_DIR / "multi_role_permission_er.png"
BUSINESS_IMAGE = ASSET_DIR / "multi_role_business_er.png"


BLUE = "#1f3fa3"
GREEN = "#1f7a36"
GRAY = "#4b5563"
LIGHT_BLUE = "#eef4ff"
LIGHT_GREEN = "#f0fbf4"
LIGHT_YELLOW = "#fff9e8"


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/Deng.ttf"),
    ]
    for font_path in candidates:
        if font_path.exists():
            return ImageFont.truetype(str(font_path), size)
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0], box[3] - box[1]


def center_text(
    draw: ImageDraw.ImageDraw,
    xy: tuple[float, float],
    text: str,
    font: ImageFont.ImageFont,
    fill: str = "black",
) -> None:
    width, height = text_size(draw, text, font)
    draw.text((xy[0] - width / 2, xy[1] - height / 2), text, font=font, fill=fill)


def draw_title(draw: ImageDraw.ImageDraw, width: int, title: str) -> None:
    title_font = load_font(46, bold=True)
    center_text(draw, (width / 2, 68), title, title_font)


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    fields: Iterable[str],
    fill: str = "white",
) -> None:
    x1, y1, x2, y2 = box
    title_font = load_font(26, bold=True)
    field_font = load_font(21)

    draw.rounded_rectangle([x1, y1, x2, y2], radius=10, fill=fill, outline=BLUE, width=3)
    draw.rectangle([x1, y1, x2, y1 + 50], fill=LIGHT_BLUE, outline=BLUE, width=3)
    center_text(draw, ((x1 + x2) / 2, y1 + 25), title, title_font)

    y = y1 + 68
    for field in fields:
        draw.text((x1 + 18, y), field, font=field_font, fill="black")
        y += 32


def draw_note(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    lines: list[str],
) -> None:
    x1, y1, x2, y2 = box
    title_font = load_font(24, bold=True)
    text_font = load_font(20)
    draw.rounded_rectangle([x1, y1, x2, y2], radius=12, fill=LIGHT_YELLOW, outline="#d6a100", width=3)
    draw.text((x1 + 18, y1 + 16), title, font=title_font, fill="#8a5a00")
    y = y1 + 58
    for line in lines:
        draw.text((x1 + 18, y), line, font=text_font, fill="black")
        y += 30


def draw_group(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    label: str,
) -> None:
    x1, y1, x2, y2 = box
    label_font = load_font(25, bold=True)
    draw.rounded_rectangle([x1, y1, x2, y2], radius=18, outline="#c9d6ee", width=4)
    draw.text((x1 + 20, y1 + 16), label, font=label_font, fill="#335f9f")


def connect(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    label: str,
    start_card: str,
    end_card: str,
    via: tuple[int, int] | None = None,
) -> None:
    font = load_font(20)
    label_font = load_font(21, bold=True)
    points = [start, via, end] if via else [start, end]
    draw.line(points, fill=GRAY, width=3)
    if via:
        lx, ly = via
    else:
        lx = (start[0] + end[0]) / 2
        ly = (start[1] + end[1]) / 2
    center_text(draw, (lx, ly - 28), label, label_font, GREEN)
    center_text(draw, (start[0], start[1] - 22), start_card, font)
    center_text(draw, (end[0], end[1] - 22), end_card, font)


def draw_permission_er(path: Path) -> None:
    width, height = 2400, 1450
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw_title(draw, width, "多角色权限 E-R 图")

    boxes = {
        "user": (130, 180, 560, 500),
        "user_role": (770, 205, 1120, 445),
        "role": (1320, 180, 1780, 560),
        "role_menu": (770, 770, 1120, 1010),
        "menu": (1320, 730, 1780, 1090),
        "lab": (130, 780, 560, 1080),
    }

    draw_group(draw, (80, 130, 1850, 1145), "用户、角色、菜单权限")

    draw_box(
        draw,
        boxes["user"],
        "用户 user",
        [
            "PK id",
            "FK laboratory_id",
            "username 登录名",
            "real_name 姓名",
            "user_no 工号/学号",
            "user_type 用户类型",
            "status 启用状态",
        ],
    )
    draw_box(
        draw,
        boxes["user_role"],
        "用户角色 user_role",
        ["PK id", "FK user_id", "FK role_id"],
        LIGHT_GREEN,
    )
    draw_box(
        draw,
        boxes["role"],
        "角色 role",
        [
            "PK id",
            "UK role_code",
            "role_name 角色名称",
            "sys_admin 系统管理员",
            "lab_director 实验室主任",
            "teacher 教师",
            "student 学生",
            "status 启用状态",
        ],
    )
    draw_box(
        draw,
        boxes["role_menu"],
        "角色菜单 role_menu",
        ["PK id", "FK role_id", "FK menu_id"],
        LIGHT_GREEN,
    )
    draw_box(
        draw,
        boxes["menu"],
        "菜单权限 menu",
        [
            "PK id",
            "FK parent_id",
            "menu_code 权限编码",
            "menu_name 菜单名称",
            "menu_type 目录/菜单/按钮",
            "path 前端路由",
            "status 启用状态",
        ],
    )
    draw_box(
        draw,
        boxes["lab"],
        "实验室 laboratory",
        [
            "PK id",
            "FK director_user_id",
            "laboratory_code",
            "laboratory_name",
            "location 位置",
            "safety_level 安全等级",
        ],
    )

    connect(draw, (560, 340), (770, 325), "分配", "1", "N")
    connect(draw, (1120, 325), (1320, 340), "对应", "N", "1")
    connect(draw, (1550, 560), (945, 770), "授权", "1", "N", (1250, 665))
    connect(draw, (1120, 890), (1320, 910), "对应", "N", "1")
    connect(draw, (240, 780), (240, 500), "包含", "1", "N")
    connect(draw, (1510, 1090), (1510, 730), "父子级", "1", "N", (1660, 910))
    connect(draw, (560, 900), (1320, 260), "负责人", "1", "1", (930, 620))

    draw_note(
        draw,
        (1920, 190, 2320, 650),
        "角色含义",
        [
            "系统管理员：用户、角色、菜单、日志",
            "实验室主任：实验室资产、审批、统计",
            "教师：借用、申领、危化品使用申请",
            "学生：授权范围内申请与个人记录",
            "实际权限由 role_menu 控制",
            "一个用户可拥有多个角色",
        ],
    )
    draw_note(
        draw,
        (1920, 760, 2320, 1085),
        "设计要点",
        [
            "user_role 解决用户与角色 M:N",
            "role_menu 解决角色与权限 M:N",
            "menu.parent_id 支持多级菜单",
            "laboratory.director_user_id 绑定主任",
            "业务表统一通过 user_id 追踪人员",
        ],
    )

    image.save(path)


def draw_business_er(path: Path) -> None:
    width, height = 2800, 1850
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw_title(draw, width, "多角色业务 E-R 图")

    boxes = {
        "role": (80, 170, 500, 520),
        "user": (650, 170, 1080, 520),
        "lab": (1230, 170, 1660, 520),
        "log": (1810, 170, 2250, 520),
        "equipment": (80, 720, 500, 1030),
        "borrow": (650, 650, 1120, 1070),
        "repair": (1260, 650, 1730, 1070),
        "calibration": (1870, 650, 2340, 1070),
        "consumable": (80, 1290, 500, 1600),
        "outbound": (650, 1230, 1120, 1660),
        "hazard": (1260, 1290, 1730, 1600),
        "haz_usage": (1870, 1230, 2340, 1660),
        "inventory": (2380, 1230, 2730, 1660),
    }

    draw_group(draw, (45, 115, 2290, 570), "公共角色与审计")
    draw_group(draw, (45, 595, 2385, 1115), "设备业务")
    draw_group(draw, (45, 1175, 2760, 1710), "耗材、危化品与库存")

    draw_box(
        draw,
        boxes["role"],
        "角色 role",
        [
            "PK id",
            "role_code",
            "sys_admin",
            "lab_director",
            "teacher",
            "student",
            "status",
        ],
    )
    draw_box(
        draw,
        boxes["user"],
        "用户 user",
        [
            "PK id",
            "FK laboratory_id",
            "username",
            "real_name",
            "user_type",
            "status",
        ],
    )
    draw_box(
        draw,
        boxes["lab"],
        "实验室 laboratory",
        ["PK id", "FK director_user_id", "code", "name", "location", "safety_level"],
    )
    draw_box(
        draw,
        boxes["log"],
        "系统日志 system_log",
        ["PK id", "FK user_id", "FK laboratory_id", "module_name", "operation_type", "operation_time"],
    )
    draw_box(
        draw,
        boxes["equipment"],
        "设备 equipment",
        ["PK id", "FK laboratory_id", "FK category_id", "equipment_code", "equipment_name", "status"],
    )
    draw_box(
        draw,
        boxes["borrow"],
        "设备借用 equipment_borrow",
        [
            "PK id",
            "FK equipment_id",
            "FK borrower_user_id",
            "FK approver_user_id",
            "borrow_time",
            "due_date",
            "return_time",
            "borrow_status",
        ],
    )
    draw_box(
        draw,
        boxes["repair"],
        "设备维修 equipment_repair",
        [
            "PK id",
            "FK equipment_id",
            "FK reporter_user_id",
            "FK repair_user_id",
            "fault_description",
            "repair_status",
            "repair_result",
        ],
    )
    draw_box(
        draw,
        boxes["calibration"],
        "设备校准 equipment_calibration",
        [
            "PK id",
            "FK equipment_id",
            "FK calibration_user_id",
            "certificate_no",
            "valid_until",
            "result",
            "status",
        ],
    )
    draw_box(
        draw,
        boxes["consumable"],
        "耗材 consumable",
        ["PK id", "FK laboratory_id", "FK category_id", "code", "name", "safe_stock", "status"],
    )
    draw_box(
        draw,
        boxes["outbound"],
        "耗材出库 consumable_outbound",
        [
            "PK id",
            "FK consumable_id",
            "FK applicant_user_id",
            "FK approver_user_id",
            "FK operator_user_id",
            "quantity",
            "outbound_status",
        ],
    )
    draw_box(
        draw,
        boxes["hazard"],
        "危化品 hazardous_material",
        ["PK id", "FK laboratory_id", "hazardous_code", "name", "CAS_no", "danger_class", "storage_location"],
    )
    draw_box(
        draw,
        boxes["haz_usage"],
        "危化品使用 hazardous_usage",
        [
            "PK id",
            "FK hazardous_material_id",
            "FK applicant_user_id",
            "FK approver_user_id",
            "FK operator_user_id",
            "witness_name",
            "action_type",
            "usage_status",
        ],
    )
    draw_box(
        draw,
        boxes["inventory"],
        "库存 inventory",
        ["PK id", "FK laboratory_id", "item_type", "item_id", "batch_no", "quantity", "warning_status"],
    )

    connect(draw, (500, 340), (650, 340), "拥有角色", "M", "N")
    connect(draw, (1080, 340), (1230, 340), "属于", "N", "1")
    connect(draw, (1080, 420), (1810, 340), "产生日志", "1", "N", (1450, 520))
    connect(draw, (500, 875), (650, 860), "被借用", "1", "N")
    connect(draw, (500, 950), (1260, 850), "被维修", "1", "N", (880, 1030))
    connect(draw, (500, 800), (1870, 850), "被校准", "1", "N", (1200, 700))
    connect(draw, (870, 520), (870, 650), "借用人: 教师/学生", "1", "N")
    connect(draw, (1010, 520), (1010, 650), "审批人: 实验室主任", "1", "N")
    connect(draw, (1430, 520), (1430, 650), "报修/维修人", "1", "N")
    connect(draw, (2050, 520), (2050, 650), "校准登记人", "1", "N")
    connect(draw, (500, 1445), (650, 1435), "出库", "1", "N")
    connect(draw, (870, 520), (870, 1230), "申领人: 教师/学生", "1", "N", (760, 880))
    connect(draw, (1010, 520), (1010, 1230), "审批/操作: 主任", "1", "N", (1120, 880))
    connect(draw, (1730, 1445), (1870, 1435), "使用", "1", "N")
    connect(draw, (1540, 520), (2050, 1230), "危化品审批", "1", "N", (1620, 900))
    connect(draw, (650, 1530), (2380, 1435), "扣减库存", "N", "1", (1500, 1740))
    connect(draw, (1730, 1530), (2380, 1530), "更新库存", "N", "1")
    connect(draw, (1660, 420), (2380, 1370), "管理库存", "1", "N", (2050, 1160))

    draw_note(
        draw,
        (2350, 180, 2720, 560),
        "角色参与规则",
        [
            "系统管理员：维护用户、角色、菜单、日志",
            "实验室主任：审批、入库、出库、校准、统计",
            "教师：设备借用、耗材申领、危化品申请",
            "学生：授权范围内提交申请与查看记录",
            "业务表用多个人员 FK 精确追踪责任",
        ],
    )

    image.save(path)


def add_paragraph(doc: Document, text: str, size: int = 12, bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)


def add_heading(doc: Document, text: str, size: int = 16) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if size >= 18 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(size)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def build_docx() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    draw_permission_er(PERMISSION_IMAGE)
    draw_business_er(BUSINESS_IMAGE)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    add_heading(doc, "实验室设备与耗材管理系统多角色详细 E-R 图", 20)
    add_paragraph(
        doc,
        "本文件重点补充多角色相关的数据关系。ER 图将系统管理员、实验室主任、教师、学生四类角色落到用户、角色、权限、业务申请、审批、操作和审计日志等数据表中，便于说明系统如何通过数据库字段追踪不同角色的责任边界。",
    )

    add_heading(doc, "一、多角色权限 E-R 图", 15)
    add_paragraph(
        doc,
        "该图描述用户、角色、菜单权限之间的核心关系。一个用户可通过 user_role 关联多个角色，一个角色可通过 role_menu 授权多个菜单或按钮权限，实验室主任通过 laboratory.director_user_id 与实验室绑定。",
    )
    doc.add_picture(str(PERMISSION_IMAGE), width=Cm(17.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "图 1 多角色权限 E-R 图")

    add_heading(doc, "二、多角色业务 E-R 图", 15)
    add_paragraph(
        doc,
        "该图描述不同角色参与核心业务时在数据表中的落点。教师和学生通常作为申请人，实验室主任作为审批人、操作人或校准登记人，系统管理员主要维护权限和审计日志。业务流水表通过 borrower_user_id、applicant_user_id、approver_user_id、operator_user_id 等外键精确记录人员责任。",
    )
    doc.add_picture(str(BUSINESS_IMAGE), width=Cm(17.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "图 2 多角色业务 E-R 图")

    add_heading(doc, "三、角色与业务字段对应说明", 15)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["角色", "主要参与业务", "对应实体/表", "关键外键字段"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
    rows = [
        ["系统管理员", "用户管理、角色权限、菜单授权、审计日志", "user、role、menu、user_role、role_menu、system_log", "user_id、role_id、menu_id"],
        ["实验室主任", "实验室资产管理、设备审批、耗材出库审批、危化品审批、校准登记", "laboratory、equipment_borrow、consumable_outbound、hazardous_usage、equipment_calibration", "director_user_id、approver_user_id、operator_user_id、calibration_user_id"],
        ["教师", "设备借用、耗材申领、危化品使用申请、设备报修", "equipment_borrow、consumable_outbound、hazardous_usage、equipment_repair", "borrower_user_id、applicant_user_id、reporter_user_id"],
        ["学生", "授权范围内借用设备、申领耗材、查看个人记录", "equipment_borrow、consumable_outbound、hazardous_usage", "borrower_user_id、applicant_user_id"],
    ]
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value

    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
    print(PERMISSION_IMAGE)
    print(BUSINESS_IMAGE)
    print(OUTPUT_DOCX)
