from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSET_DIR = DOCS_DIR / "generated_assets"
OUTPUT_IMAGE = ASSET_DIR / "reference_style_role_usecase_flow.png"
OUTPUT_DOCX = DOCS_DIR / "reference_style_role_usecase_flow.docx"

BLUE = "#3f6ecb"
BLUE_DARK = "#2f5fb3"
TEXT = "#214a9a"
WARN_BG = "#fff1c4"
WARN_TEXT = "#b56b00"
MUTED = "#666666"


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/Deng.ttf"),
    ]
    for font_path in candidates:
        if font_path.exists():
            return ImageFont.truetype(str(font_path), size)
    return ImageFont.load_default()


def center_text(draw: ImageDraw.ImageDraw, xy: tuple[float, float], text: str, font: ImageFont.ImageFont, fill: str) -> None:
    box = draw.textbbox((0, 0), text, font=font)
    width = box[2] - box[0]
    height = box[3] - box[1]
    draw.text((xy[0] - width / 2, xy[1] - height / 2), text, font=font, fill=fill)


def draw_actor(draw: ImageDraw.ImageDraw, x: int, y: int, label: str, scale: float = 1.0) -> None:
    font = load_font(int(24 * scale), bold=True)
    head = int(18 * scale)
    body_w = int(56 * scale)
    body_h = int(84 * scale)
    arm = int(30 * scale)
    leg_w = int(15 * scale)
    leg_h = int(48 * scale)

    draw.ellipse([x - head, y - 70 * scale, x + head, y - 34 * scale], fill=BLUE_DARK)
    draw.rounded_rectangle([x - body_w / 2, y - 32 * scale, x + body_w / 2, y + body_h / 2], radius=6, fill=BLUE_DARK)
    draw.rectangle([x - body_w / 2 - arm, y - 20 * scale, x - body_w / 2, y + 28 * scale], fill=BLUE_DARK)
    draw.rectangle([x + body_w / 2, y - 20 * scale, x + body_w / 2 + arm, y + 28 * scale], fill=BLUE_DARK)
    draw.rectangle([x - 18 * scale, y + body_h / 2, x - 18 * scale + leg_w, y + body_h / 2 + leg_h], fill=BLUE_DARK)
    draw.rectangle([x + 4 * scale, y + body_h / 2, x + 4 * scale + leg_w, y + body_h / 2 + leg_h], fill=BLUE_DARK)
    center_text(draw, (x, y + 120 * scale), label, font, TEXT)


def draw_usecase(draw: ImageDraw.ImageDraw, x: int, y: int, text: str, w: int = 220, h: int = 64) -> None:
    font = load_font(24, bold=True)
    draw.ellipse([x - w / 2, y - h / 2, x + w / 2, y + h / 2], fill=BLUE, outline=BLUE_DARK, width=2)
    center_text(draw, (x, y), text, font, "white")


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    label: str | None = None,
    width: int = 2,
) -> None:
    sx, sy = start
    ex, ey = end
    draw.line([start, end], fill=BLUE_DARK, width=width)
    dx = ex - sx
    dy = ey - sy
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux = dx / length
    uy = dy / length
    size = 12
    left = (ex - ux * size - uy * size * 0.55, ey - uy * size + ux * size * 0.55)
    right = (ex - ux * size + uy * size * 0.55, ey - uy * size - ux * size * 0.55)
    draw.polygon([end, left, right], fill=BLUE_DARK)
    if label:
        font = load_font(17)
        center_text(draw, ((sx + ex) / 2, (sy + ey) / 2 - 15), label, font, BLUE_DARK)


def draw_warning(draw: ImageDraw.ImageDraw, width: int) -> None:
    font = load_font(20)
    draw.rectangle([150, 18, width - 150, 62], fill=WARN_BG)
    draw.text((168, 30), "△ 多角色用例流程图需覆盖所有核心功能，画完后应对照功能清单检查无遗漏。", font=font, fill=WARN_TEXT)


def draw_diagram(path: Path) -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1720, 2900
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(40, bold=True)
    note_font = load_font(18)

    draw_warning(draw, width)
    center_text(draw, (width / 2, 105), "实验室设备与耗材管理系统多角色用例流程说明图", title_font, TEXT)

    actors = {
        "学生": (175, 330, 1.0),
        "教师": (175, 675, 1.0),
        "实验室主任": (1460, 690, 1.0),
        "设备管理员": (1440, 1110, 0.86),
        "耗材管理员": (1440, 1390, 0.86),
        "危化品管理员": (1440, 1670, 0.86),
        "维修人员": (1440, 1940, 0.82),
        "校准人员": (1440, 2180, 0.82),
        "系统管理员": (1460, 2580, 0.9),
    }
    for label, (x, y, scale) in actors.items():
        draw_actor(draw, x, y, label, scale)

    usecases = {
        # Common and applicant side.
        "登录认证": (520, 170, 205),
        "个人中心": (520, 280, 205),
        "查看个人记录": (520, 390, 230),
        "设备借用申请": (520, 545, 250),
        "耗材申领申请": (520, 665, 250),
        "危化品使用申请": (520, 785, 270),
        "设备报修": (520, 905, 205),
        # Workflow center.
        "审批处理": (865, 555, 220),
        "设备归还登记": (865, 675, 250),
        "库存数量校验": (865, 795, 260),
        "双人复核": (865, 915, 220),
        "废液处理": (865, 1035, 220),
        "到期提醒": (865, 1155, 220),
        # Equipment admin.
        "设备分类管理": (560, 1080, 250),
        "设备台账维护": (560, 1200, 250),
        "借用登记确认": (560, 1320, 250),
        "归还验收": (560, 1440, 220),
        "维修派单": (560, 1560, 220),
        "逾期催还": (560, 1680, 220),
        # Consumable and hazardous admins.
        "耗材分类管理": (895, 1290, 250),
        "耗材台账维护": (895, 1410, 250),
        "耗材入库": (895, 1530, 220),
        "出库确认": (895, 1650, 220),
        "批次库存维护": (895, 1770, 250),
        "安全库存预警": (895, 1890, 250),
        "危化品台账": (895, 2035, 230),
        "CAS/MSDS维护": (895, 2155, 250),
        "领用登记": (895, 2275, 220),
        "余量归还": (895, 2395, 220),
        "安全审计": (895, 2515, 220),
        "废液登记": (1180, 2275, 220),
        # Maintenance/calibration/admin.
        "接收维修任务": (560, 1880, 250),
        "记录维修过程": (560, 2000, 250),
        "登记维修费用": (560, 2120, 250),
        "提交维修结果": (560, 2240, 250),
        "接收校准任务": (1180, 1880, 250),
        "登记证书编号": (1180, 2000, 250),
        "记录校准结果": (1180, 2120, 250),
        "更新下次校准": (1180, 2240, 250),
        "用户管理": (520, 2440, 220),
        "角色权限管理": (520, 2560, 250),
        "菜单权限管理": (520, 2680, 250),
        "审计日志": (1180, 2500, 220),
        "系统配置": (1180, 2620, 220),
    }
    for label, (x, y, w) in usecases.items():
        draw_usecase(draw, x, y, label, w=w)

    def connect_actor(actor: str, targets: list[str], from_shift: tuple[int, int] = (62, 0), to_left: bool = True) -> None:
        ax, ay, _ = actors[actor]
        for target in targets:
            tx, ty, tw = usecases[target]
            if to_left:
                arrow(draw, (ax + from_shift[0], ay + from_shift[1]), (int(tx - tw / 2), ty))
            else:
                arrow(draw, (ax - from_shift[0], ay + from_shift[1]), (int(tx + tw / 2), ty))

    connect_actor("学生", ["登录认证", "个人中心", "查看个人记录", "设备借用申请", "耗材申领申请", "危化品使用申请"])
    connect_actor("教师", ["登录认证", "个人中心", "查看个人记录", "设备借用申请", "耗材申领申请", "危化品使用申请", "设备报修"])
    connect_actor(
        "实验室主任",
        ["审批处理", "设备归还登记", "库存数量校验", "双人复核", "废液处理", "到期提醒", "审计日志"],
        from_shift=(70, 0),
        to_left=False,
    )
    connect_actor(
        "设备管理员",
        ["设备分类管理", "设备台账维护", "借用登记确认", "归还验收", "维修派单", "逾期催还"],
        from_shift=(62, 0),
        to_left=False,
    )
    connect_actor(
        "耗材管理员",
        ["耗材分类管理", "耗材台账维护", "耗材入库", "出库确认", "批次库存维护", "安全库存预警"],
        from_shift=(62, 0),
        to_left=False,
    )
    connect_actor(
        "危化品管理员",
        ["危化品台账", "CAS/MSDS维护", "领用登记", "余量归还", "废液登记", "安全审计"],
        from_shift=(62, 0),
        to_left=False,
    )
    connect_actor("维修人员", ["接收维修任务", "记录维修过程", "登记维修费用", "提交维修结果"], from_shift=(62, 0), to_left=False)
    connect_actor("校准人员", ["接收校准任务", "登记证书编号", "记录校准结果", "更新下次校准"], from_shift=(62, 0), to_left=False)
    connect_actor("系统管理员", ["用户管理", "角色权限管理", "菜单权限管理", "审计日志", "系统配置"], from_shift=(70, 0), to_left=False)

    # Include/extend style links between use cases, close to the sample image.
    arrow(draw, (640, 545), (755, 555), "包含")
    arrow(draw, (640, 665), (745, 795), "包含")
    arrow(draw, (650, 785), (745, 795), "包含")
    arrow(draw, (995, 795), (995, 915), "扩展")
    arrow(draw, (995, 915), (995, 1035), "扩展")
    arrow(draw, (995, 1155), (1020, 1890), "扩展")
    arrow(draw, (670, 905), (435, 1880), "扩展")
    arrow(draw, (670, 1560), (435, 1880), "派单")
    arrow(draw, (680, 1680), (760, 1155), "扩展")
    arrow(draw, (1020, 2515), (1010, 2410), "扩展")

    # Supervision/inheritance lines on the right.
    arrow(draw, (1460, 2480), (1460, 2020), "监管")
    arrow(draw, (1460, 2480), (1460, 1750), "监管")
    arrow(draw, (1460, 2480), (1460, 1470), "监管")
    arrow(draw, (1460, 2480), (1460, 1190), "监管")
    arrow(draw, (1460, 2480), (1460, 780), "继承/监管")

    draw.text(
        (135, 2825),
        "说明：学生/教师主要发起申请；实验室主任负责审批、复核和验收；设备、耗材、危化品管理员负责专项台账与库存；维修/校准人员负责专业处理；系统管理员负责权限、配置和审计。",
        font=note_font,
        fill=MUTED,
    )

    image.save(path)


def build_docx() -> None:
    draw_diagram(OUTPUT_IMAGE)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("实验室设备与耗材管理系统多角色用例流程说明图")
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(18)

    doc.add_picture(str(OUTPUT_IMAGE), width=Cm(17.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = caption.add_run("图 1 多角色用例流程说明图")
    cap.font.name = "Arial"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cap.font.size = Pt(11)

    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
    print(OUTPUT_IMAGE)
    print(OUTPUT_DOCX)
