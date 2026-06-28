from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSET_DIR = DOCS_DIR / "generated_assets"
SOURCE_DOCX = DOCS_DIR / "lab_management_design_report.docx"
OUTPUT_DOCX = DOCS_DIR / "lab_management_design_report_no_code.docx"
OUTPUT_CN_DOCX = DOCS_DIR / "实验室设备与耗材管理系统设计说明书_无代码版.docx"
OBJECT_IMAGE = ASSET_DIR / "database_object_design_overview.png"
FLOW_IMAGE = ASSET_DIR / "database_process_flow_no_code.png"


BLUE = "#3f6ecb"
BLUE_DARK = "#2f5fb3"
GREEN = "#2f8a4c"
ORANGE = "#e48a23"
LIGHT_BLUE = "#eef4ff"
LIGHT_GREEN = "#effaf2"
LIGHT_ORANGE = "#fff4e7"
TEXT = "#1d3557"


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/Deng.ttf"),
    ]
    for font_path in candidates:
        if font_path.exists():
            return ImageFont.truetype(str(font_path), size)
    return ImageFont.load_default()


def center_text(draw: ImageDraw.ImageDraw, xy: tuple[float, float], text: str, font: ImageFont.ImageFont, fill: str = "black") -> None:
    box = draw.textbbox((0, 0), text, font=font)
    draw.text((xy[0] - (box[2] - box[0]) / 2, xy[1] - (box[3] - box[1]) / 2), text, font=font, fill=fill)


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    lines: list[str],
    fill: str,
    outline: str,
) -> None:
    x1, y1, x2, y2 = box
    title_font = load_font(25, bold=True)
    text_font = load_font(20)
    draw.rounded_rectangle([x1, y1, x2, y2], radius=16, fill=fill, outline=outline, width=3)
    center_text(draw, ((x1 + x2) / 2, y1 + 34), title, title_font, outline)
    y = y1 + 76
    for line in lines:
        draw.text((x1 + 24, y), line, font=text_font, fill=TEXT)
        y += 34


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], label: str | None = None) -> None:
    draw.line([start, end], fill=BLUE_DARK, width=3)
    sx, sy = start
    ex, ey = end
    dx = ex - sx
    dy = ey - sy
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux = dx / length
    uy = dy / length
    size = 14
    left = (ex - ux * size - uy * size * 0.6, ey - uy * size + ux * size * 0.6)
    right = (ex - ux * size + uy * size * 0.6, ey - uy * size - ux * size * 0.6)
    draw.polygon([end, left, right], fill=BLUE_DARK)
    if label:
        center_text(draw, ((sx + ex) / 2, (sy + ey) / 2 - 22), label, load_font(18), BLUE_DARK)


def draw_object_overview(path: Path) -> None:
    width, height = 1800, 1100
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(42, bold=True)
    center_text(draw, (width / 2, 62), "数据库物理设计与对象设计总览图", title_font, TEXT)

    draw_box(
        draw,
        (80, 140, 520, 470),
        "物理索引设计",
        [
            "设备编码唯一索引",
            "设备状态普通索引",
            "借用状态 + 应还日期复合索引",
            "库存物资 + 批次复合索引",
            "预警状态 + 有效期复合索引",
            "危化品追溯复合索引",
        ],
        LIGHT_BLUE,
        BLUE_DARK,
    )
    draw_box(
        draw,
        (680, 140, 1120, 470),
        "视图设计",
        [
            "设备借用状态视图",
            "整合设备、实验室、借用人、审批人",
            "用于借用列表、报表中心",
            "隐藏复杂联表逻辑",
            "统一前后端展示口径",
        ],
        LIGHT_GREEN,
        GREEN,
    )
    draw_box(
        draw,
        (1280, 140, 1720, 470),
        "存储过程设计",
        [
            "耗材出库过程",
            "统一校验批次库存",
            "库存不足时终止出库",
            "出库成功后扣减库存",
            "减少业务入口重复实现",
        ],
        LIGHT_ORANGE,
        ORANGE,
    )
    draw_box(
        draw,
        (380, 620, 820, 940),
        "触发器设计",
        [
            "设备借用后自动更新状态",
            "归还后恢复可用状态",
            "耗材入库后维护库存批次",
            "危化品使用后更新余量",
            "关键状态变化写入日志",
        ],
        LIGHT_ORANGE,
        ORANGE,
    )
    draw_box(
        draw,
        (980, 620, 1420, 940),
        "查询统计设计",
        [
            "按用户和时间查询借用记录",
            "统计本月耗材出库数量",
            "统计本月耗材出库金额",
            "支持报表中心筛选导出",
            "辅助主任和管理员决策",
        ],
        LIGHT_BLUE,
        BLUE_DARK,
    )

    arrow(draw, (520, 300), (680, 300), "加速查询")
    arrow(draw, (1120, 300), (1280, 300), "服务业务")
    arrow(draw, (1500, 470), (1220, 620), "联动库存")
    arrow(draw, (820, 790), (980, 790), "沉淀统计")
    arrow(draw, (680, 620), (900, 470), "保障一致性")

    image.save(path)


def draw_process_flow(path: Path) -> None:
    width, height = 1900, 1250
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(42, bold=True)
    center_text(draw, (width / 2, 62), "数据库对象业务流程说明图", title_font, TEXT)

    steps1 = [
        ("出库申请", "申请人提交耗材出库需求"),
        ("库存校验", "检查物资、批次和数量"),
        ("审批确认", "主任或管理员确认出库"),
        ("扣减库存", "更新对应批次数量"),
        ("预警判断", "低库存或临期生成提醒"),
    ]
    steps2 = [
        ("借用申请", "教师/学生提交设备借用"),
        ("状态校验", "确认设备处于可用状态"),
        ("审批通过", "实验室主任审核"),
        ("状态更新", "设备自动变为借出"),
        ("归还登记", "验收后恢复可用并写日志"),
    ]

    def draw_steps(y: int, title: str, steps: list[tuple[str, str]]) -> None:
        title_f = load_font(28, bold=True)
        center_text(draw, (150, y), title, title_f, TEXT)
        x = 330
        for index, (name, desc) in enumerate(steps):
            draw_box(draw, (x, y - 70, x + 250, y + 80), name, [desc], LIGHT_BLUE if index % 2 == 0 else LIGHT_GREEN, BLUE_DARK if index % 2 == 0 else GREEN)
            if index < len(steps) - 1:
                arrow(draw, (x + 250, y + 5), (x + 330, y + 5))
            x += 330

    draw_steps(330, "耗材出库", steps1)
    draw_steps(790, "设备借用", steps2)

    note_font = load_font(22)
    draw.rounded_rectangle([120, 1040, 1780, 1160], radius=18, fill="#fffaf0", outline=ORANGE, width=3)
    draw.text(
        (160, 1078),
        "说明：本图用流程方式替代 SQL 代码。视图负责统一查询口径，存储过程负责库存校验和扣减，触发器负责设备状态与库存状态自动联动。",
        font=note_font,
        fill=TEXT,
    )

    image.save(path)


def is_code_paragraph(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    prefixes = (
        "CREATE ",
        "SELECT ",
        "UPDATE ",
        "DELIMITER",
        "BEGIN",
        "END",
        "IF ",
        "WHERE ",
        "FROM ",
        "LEFT JOIN",
        "SIGNAL ",
        "SET ",
        "AND ",
        "IN ",
    )
    return stripped.startswith(prefixes) or stripped in {")", "END //", "DELIMITER ;"}


def add_image_after_heading(doc: Document, heading_text: str, image_path: Path, caption: str) -> None:
    target = None
    for paragraph in doc.paragraphs:
        if heading_text in paragraph.text:
            target = paragraph
            break
    if target is None:
        return

    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(image_path), width=Cm(16.2))

    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)

    for p in [image_paragraph, caption_paragraph]:
        p._p.getparent().remove(p._p)

    target._p.addnext(caption_paragraph._p)
    target._p.addnext(image_paragraph._p)


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    draw_object_overview(OBJECT_IMAGE)
    draw_process_flow(FLOW_IMAGE)

    doc = Document(SOURCE_DOCX)

    for paragraph in list(doc.paragraphs):
        if is_code_paragraph(paragraph.text):
            paragraph._element.getparent().remove(paragraph._element)

    add_image_after_heading(doc, "2.3  数据库物理设计", OBJECT_IMAGE, "图3 数据库物理设计与对象设计总览图")
    add_image_after_heading(doc, "2.4  SQL Server数据库对象设计", FLOW_IMAGE, "图4 数据库对象业务流程说明图")

    doc.save(OUTPUT_DOCX)
    doc.save(OUTPUT_CN_DOCX)
    print(OUTPUT_DOCX)
    print(OUTPUT_CN_DOCX)
    print(OBJECT_IMAGE)
    print(FLOW_IMAGE)


if __name__ == "__main__":
    main()
