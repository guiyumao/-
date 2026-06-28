from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSET_DIR = DOCS_DIR / "generated_assets"
OUTPUT_IMAGE = ASSET_DIR / "multi_role_usecase_flow.png"
OUTPUT_DOCX = DOCS_DIR / "multi_role_usecase_flow.docx"

BLUE = "#3f6ecb"
LINE = "#2f5fb3"
TEXT = "#1d3f8f"
WARN_BG = "#fff1c4"
WARN_TEXT = "#b56b00"


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/Deng.ttf"),
    ]
    for font_path in candidates:
        if font_path.exists():
            return ImageFont.truetype(str(font_path), size)
    return ImageFont.load_default()


def center_text(draw: ImageDraw.ImageDraw, xy: tuple[float, float], text: str, font: ImageFont.ImageFont, fill: str) -> None:
    box = draw.textbbox((0, 0), text, font=font)
    width = box[2] - box[0]
    height = box[3] - box[1]
    draw.text((xy[0] - width / 2, xy[1] - height / 2), text, font=font, fill=fill)


def draw_actor(draw: ImageDraw.ImageDraw, x: int, y: int, label: str) -> None:
    label_font = load_font(24)
    draw.ellipse([x - 16, y - 62, x + 16, y - 30], fill=LINE)
    draw.rounded_rectangle([x - 24, y - 28, x + 24, y + 50], radius=6, fill=LINE)
    draw.rectangle([x - 50, y - 20, x - 24, y + 30], fill=LINE)
    draw.rectangle([x + 24, y - 20, x + 50, y + 30], fill=LINE)
    draw.rectangle([x - 18, y + 50, x - 4, y + 95], fill=LINE)
    draw.rectangle([x + 4, y + 50, x + 18, y + 95], fill=LINE)
    center_text(draw, (x, y + 130), label, label_font, TEXT)


def draw_usecase(draw: ImageDraw.ImageDraw, x: int, y: int, text: str, w: int = 190, h: int = 66) -> None:
    font = load_font(25, bold=True)
    draw.ellipse([x - w / 2, y - h / 2, x + w / 2, y + h / 2], fill=BLUE, outline=LINE, width=2)
    center_text(draw, (x, y), text, font, "white")


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], dashed: bool = False, label: str | None = None) -> None:
    if dashed:
        draw.line([start, end], fill=LINE, width=2)
        # Pillow has no portable dashed line helper; label distinguishes extend/include links.
    else:
        draw.line([start, end], fill=LINE, width=2)

    sx, sy = start
    ex, ey = end
    dx = ex - sx
    dy = ey - sy
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux = dx / length
    uy = dy / length
    size = 12
    left = (ex - ux * size - uy * size * 0.55, ey - uy * size + ux * size * 0.55)
    right = (ex - ux * size + uy * size * 0.55, ey - uy * size - ux * size * 0.55)
    draw.polygon([end, left, right], fill=LINE)

    if label:
        font = load_font(18)
        mx = (sx + ex) / 2
        my = (sy + ey) / 2 - 14
        center_text(draw, (mx, my), label, font, LINE)


def draw_warning(draw: ImageDraw.ImageDraw) -> None:
    font = load_font(22)
    draw.rectangle([130, 10, 1220, 56], fill=WARN_BG)
    draw.text((150, 22), "△ 多角色用例流程图需覆盖核心业务，画完后应对照功能清单检查无遗漏。", font=font, fill=WARN_TEXT)


def draw_diagram(path: Path) -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1500, 1760
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(42, bold=True)
    note_font = load_font(20)

    draw_warning(draw)
    center_text(draw, (width / 2, 92), "实验室设备与耗材管理系统多角色用例流程说明图", title_font, TEXT)

    actors = {
        "学生": (160, 350),
        "教师": (160, 720),
        "实验室主任": (1260, 780),
        "系统管理员": (1260, 1280),
    }
    for label, (x, y) in actors.items():
        draw_actor(draw, x, y, label)

    usecases = {
        "登录认证": (470, 170),
        "个人中心": (470, 270),
        "查看个人记录": (470, 370),
        "设备借用申请": (470, 500),
        "耗材申领": (470, 610),
        "危化品使用申请": (470, 720),
        "设备报修": (470, 830),
        "设备归还登记": (780, 500),
        "审批处理": (780, 610),
        "库存数量校验": (780, 720),
        "双人复核": (780, 830),
        "废液处理": (780, 940),
        "实验室管理": (780, 1080),
        "设备管理": (780, 1190),
        "耗材管理": (780, 1300),
        "危化品管理": (780, 1410),
        "库存预警": (780, 1520),
        "用户管理": (470, 1120),
        "角色权限管理": (470, 1230),
        "菜单权限管理": (470, 1340),
        "审计日志": (470, 1450),
        "报表统计": (470, 1560),
    }

    for text, (x, y) in usecases.items():
        draw_usecase(draw, x, y, text, w=220 if len(text) >= 6 else 190)

    # Student and teacher common application flows.
    for actor in ["学生", "教师"]:
        ax, ay = actors[actor]
        for target in ["登录认证", "个人中心", "查看个人记录"]:
            x, y = usecases[target]
            arrow(draw, (ax + 55, ay - 30), (x - 110, y), label=None)
        for target in ["设备借用申请", "耗材申领", "危化品使用申请"]:
            x, y = usecases[target]
            arrow(draw, (ax + 55, ay + 10), (x - 110, y), label=None)

    # Teacher-specific repair flow.
    ax, ay = actors["教师"]
    x, y = usecases["设备报修"]
    arrow(draw, (ax + 55, ay + 40), (x - 110, y), label=None)

    # Director business and management flows.
    ax, ay = actors["实验室主任"]
    for target in [
        "设备归还登记",
        "审批处理",
        "库存数量校验",
        "双人复核",
        "废液处理",
        "实验室管理",
        "设备管理",
        "耗材管理",
        "危化品管理",
        "库存预警",
        "报表统计",
    ]:
        x, y = usecases[target]
        arrow(draw, (ax - 60, ay + 20), (x + 110, y), label=None)

    # Admin flows.
    ax, ay = actors["系统管理员"]
    for target in ["用户管理", "角色权限管理", "菜单权限管理", "审计日志", "报表统计"]:
        x, y = usecases[target]
        arrow(draw, (ax - 60, ay + 10), (x + 110, y), label=None)

    # Role generalization line: system admin inherits director-level management in this system.
    arrow(draw, (1260, 1210), (1260, 950), label="继承/监管")

    # Include/extend relationships.
    arrow(draw, (580, 500), (670, 610), dashed=True, label="包含")
    arrow(draw, (580, 610), (670, 720), dashed=True, label="包含")
    arrow(draw, (580, 720), (670, 720), dashed=True, label="包含")
    arrow(draw, (890, 720), (890, 830), dashed=True, label="扩展")
    arrow(draw, (890, 830), (890, 940), dashed=True, label="扩展")
    arrow(draw, (890, 1520), (580, 1560), dashed=True, label="扩展")
    arrow(draw, (580, 1450), (670, 610), dashed=True, label="扩展")

    # Small explanatory note.
    draw.text(
        (130, 1685),
        "说明：学生和教师主要发起申请；实验室主任负责审批、复核、库存和资产管理；系统管理员负责用户、角色、菜单权限和审计日志。",
        font=note_font,
        fill="#555555",
    )

    image.save(path)


def build_docx() -> None:
    draw_diagram(OUTPUT_IMAGE)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("实验室设备与耗材管理系统多角色用例流程说明图")
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(18)

    doc.add_picture(str(OUTPUT_IMAGE), width=Cm(17.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = caption.add_run("图 1 多角色用例流程说明图")
    cap.font.name = "Arial"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cap.font.size = Pt(11)

    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
    print(OUTPUT_IMAGE)
    print(OUTPUT_DOCX)
