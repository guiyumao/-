from __future__ import annotations

import math
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import matplotlib.pyplot as plt
from matplotlib.patches import Ellipse, Polygon, Rectangle
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
SOURCE_MD = DOCS_DIR / "database_and_requirements_analysis.md"
OUTPUT_DOCX = DOCS_DIR / "实验室设备与耗材管理系统需求分析.docx"
ASSET_DIR = DOCS_DIR / "generated_assets"


ENTITY_ATTRIBUTES: dict[str, list[str]] = {
    "用户(User)": ["用户编号", "所属实验室", "用户名", "密码", "姓名", "手机号", "邮箱", "状态"],
    "角色(Role)": ["角色编号", "角色编码", "角色名称", "状态", "备注"],
    "菜单(Menu)": ["菜单编号", "父菜单", "菜单名称", "路由", "权限标识", "类型", "状态"],
    "实验室(Laboratory)": ["实验室编号", "实验室编码", "名称", "位置", "安全等级", "负责人", "状态"],
    "设备类别(EquipmentCategory)": ["类别编号", "类别编码", "类别名称", "排序号"],
    "设备(Equipment)": ["设备编号", "所属实验室", "设备类别", "名称", "型号", "品牌", "状态", "校准周期"],
    "耗材类别(ConsumableCategory)": ["类别编号", "类别编码", "类别名称", "排序号"],
    "耗材(Consumable)": ["耗材编号", "所属实验室", "耗材类别", "名称", "规格", "单位", "安全库存", "状态"],
    "危化品(HazardousMaterial)": ["危化品编号", "所属实验室", "名称", "CAS号", "危险类别", "浓度/纯度", "储位", "状态"],
    "设备借用(EquipmentBorrow)": ["借用记录编号", "设备", "借用人", "审批人", "借用时间", "应还时间", "归还时间", "状态"],
    "设备维修(EquipmentRepair)": ["维修记录编号", "设备", "报修人", "维修人", "报修时间", "维修费用", "维修结果", "状态"],
    "设备校准(EquipmentCalibration)": ["校准记录编号", "设备", "校准人", "校准日期", "有效期至", "证书编号", "结果", "状态"],
    "耗材入库(ConsumableInbound)": ["入库记录编号", "耗材", "操作人", "批次号", "数量", "单价", "供应商", "入库日期"],
    "耗材出库(ConsumableOutbound)": ["出库记录编号", "耗材", "申请人", "审批人", "操作人", "批次号", "数量", "出库日期"],
    "危化品使用(HazardousUsage)": ["使用记录编号", "危化品", "申请人", "审批人", "操作人", "批次号", "数量", "业务类型"],
    "库存(Inventory)": ["库存编号", "所属实验室", "物资类型", "物资编号", "批次号", "库存数量", "锁定数量", "预警状态"],
    "系统日志(SystemLog)": ["日志编号", "操作用户", "所属实验室", "模块名称", "操作类型", "业务主键", "操作时间", "状态"],
}


RELATIONS = [
    ("用户", "拥有", "角色", "M:N"),
    ("角色", "授权", "菜单", "M:N"),
    ("菜单", "父子", "菜单", "1:N"),
    ("实验室", "包含", "用户", "1:N"),
    ("实验室", "拥有", "设备", "1:N"),
    ("实验室", "拥有", "耗材", "1:N"),
    ("实验室", "拥有", "危化品", "1:N"),
    ("实验室", "管理", "库存", "1:N"),
    ("设备类别", "分类", "设备", "1:N"),
    ("耗材类别", "分类", "耗材", "1:N"),
    ("设备", "借用", "设备借用", "1:N"),
    ("设备", "维修", "设备维修", "1:N"),
    ("设备", "校准", "设备校准", "1:N"),
    ("用户", "申请", "设备借用", "1:N"),
    ("用户", "报修", "设备维修", "1:N"),
    ("用户", "执行", "设备校准", "1:N"),
    ("耗材", "入库", "耗材入库", "1:N"),
    ("耗材", "出库", "耗材出库", "1:N"),
    ("危化品", "使用", "危化品使用", "1:N"),
    ("耗材", "形成批次", "库存", "1:N"),
    ("危化品", "形成批次", "库存", "1:N"),
    ("用户", "产生", "系统日志", "1:N"),
]


@dataclass
class SectionBlock:
    title: str
    level: int
    body: list[str]


def ensure_assets_dir() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def parse_markdown_sections(text: str) -> list[SectionBlock]:
    lines = text.splitlines()
    sections: list[SectionBlock] = []
    current: SectionBlock | None = None

    for line in lines:
        m = re.match(r"^(#{1,3})\s+(.*)$", line)
        if m:
            if current:
                sections.append(current)
            current = SectionBlock(title=m.group(2).strip(), level=len(m.group(1)), body=[])
        elif current is not None:
            current.body.append(line.rstrip())

    if current:
        sections.append(current)
    return sections


def get_section_map(sections: Iterable[SectionBlock]) -> dict[str, SectionBlock]:
    return {section.title: section for section in sections}


def clean_title(title: str) -> str:
    return re.sub(r"（.*?）", "", title).strip()


def get_body_text(section: SectionBlock | None) -> list[str]:
    if not section:
        return []
    return [line for line in section.body if line.strip()]


def normalize_table(block_lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in block_lines:
        if not line.strip().startswith("|"):
            continue
        if set(line.replace("|", "").replace("-", "").replace(" ", "")) == set():
            continue
        cols = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if cols and not all(set(col) <= {"-"} for col in cols):
            rows.append(cols)
    return rows


def split_mixed_content(lines: list[str]) -> tuple[list[str], list[list[list[str]]]]:
    paragraphs: list[str] = []
    tables: list[list[list[str]]] = []
    current_table_lines: list[str] = []

    def flush_table() -> None:
        nonlocal current_table_lines
        if current_table_lines:
            table = normalize_table(current_table_lines)
            if table:
                tables.append(table)
            current_table_lines = []

    for line in lines:
        if line.strip().startswith("|"):
            current_table_lines.append(line)
        else:
            flush_table()
            if line.strip().startswith("```"):
                continue
            if line.strip():
                paragraphs.append(line.strip())
    flush_table()
    return paragraphs, tables


def set_document_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def set_cell_text(cell, text: str, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            set_cell_text(table.cell(i, j), value, bold=(i == 0), center=True)
    doc.add_paragraph("")


def add_paragraph(doc: Document, text: str, first_line_indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)


def add_bullet_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f"（1）{text}" if not re.match(r"^（\d+）", text) else text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_cover_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("AI 时代软件开发课程设计")
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(22)
    run.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("个人设计文档")
    run2.font.name = "黑体"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run2.font.size = Pt(20)
    run2.bold = True

    doc.add_paragraph("")
    cover_table = doc.add_table(rows=7, cols=2)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_table.style = "Table Grid"
    labels = [
        ("课程名称", "数据库原理与应用"),
        ("系统名称", "实验室设备与耗材管理系统"),
        ("姓名", "卫德翔"),
        ("学号", "202400406103"),
        ("班级", "软件243"),
        ("指导教师", ""),
        ("提交日期", "2026年6月15日"),
    ]
    for i, (left, right) in enumerate(labels):
        set_cell_text(cover_table.cell(i, 0), left, bold=True, center=True)
        set_cell_text(cover_table.cell(i, 1), right, center=True)

    doc.add_paragraph("")
    add_paragraph(
        doc,
        "文档说明：本文档参照课程设计模板组织需求分析内容，围绕实验室设备与耗材管理系统的业务需求、用例分析、概念结构设计与数据库设计要求进行展开，并重点补充采用经典 Chen 模型绘制的系统总体 E-R 图与实体属性图。",
        first_line_indent=False,
    )
    doc.add_page_break()


def add_doc_intro(doc: Document) -> None:
    add_heading(doc, "STEP 01 需求分析", 1)
    add_paragraph(
        doc,
        "本阶段围绕系统边界、角色职责、业务流程与数据库概念结构展开分析，目标是为后续数据库逻辑设计、框架搭建和模块开发提供一致、清晰且可落地的设计基础。",
    )


def add_section_from_md(doc: Document, title: str, lines: list[str]) -> None:
    paragraphs, tables = split_mixed_content(lines)
    for text in paragraphs:
        if re.match(r"^\d+\.", text):
            add_paragraph(doc, text)
        elif text.startswith("- "):
            add_paragraph(doc, text[2:], first_line_indent=False)
        else:
            add_paragraph(doc, text)
    for table in tables:
        add_table(doc, table)


def draw_system_er(path: Path) -> None:
    plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False

    fig, ax = plt.subplots(figsize=(16, 10))
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis("off")

    entities = {
        "用户": (12, 82),
        "角色": (30, 82),
        "菜单": (48, 82),
        "实验室": (12, 58),
        "设备类别": (30, 58),
        "设备": (48, 58),
        "耗材类别": (66, 58),
        "耗材": (84, 58),
        "危化品": (84, 36),
        "库存": (66, 36),
        "设备借用": (48, 36),
        "设备维修": (30, 36),
        "设备校准": (12, 36),
        "耗材入库": (66, 18),
        "耗材出库": (84, 18),
        "危化品使用": (48, 18),
        "系统日志": (12, 18),
    }

    def rect(center_x: float, center_y: float, text: str) -> None:
        w, h = 12, 6
        ax.add_patch(Rectangle((center_x - w / 2, center_y - h / 2), w, h, fill=False, linewidth=1.5))
        ax.text(center_x, center_y, text, ha="center", va="center", fontsize=11)

    def diamond(center_x: float, center_y: float, text: str) -> None:
        w, h = 8, 5
        points = [
            (center_x, center_y + h / 2),
            (center_x + w / 2, center_y),
            (center_x, center_y - h / 2),
            (center_x - w / 2, center_y),
        ]
        ax.add_patch(Polygon(points, closed=True, fill=False, linewidth=1.2))
        ax.text(center_x, center_y, text, ha="center", va="center", fontsize=9)

    def connect(a: str, rel: str, b: str, card: str, dx: float = 0, dy: float = 0) -> None:
        x1, y1 = entities[a]
        x2, y2 = entities[b]
        mx = (x1 + x2) / 2 + dx
        my = (y1 + y2) / 2 + dy
        diamond(mx, my, rel)
        ax.plot([x1, mx - math.copysign(4, mx - x1 if mx != x1 else 1)], [y1, my], color="black", linewidth=1)
        ax.plot([mx + math.copysign(4, x2 - mx if x2 != mx else 1), x2], [my, y2], color="black", linewidth=1)
        ax.text((x1 + mx) / 2, (y1 + my) / 2 + 1.5, card.split(":")[0], fontsize=8, ha="center")
        ax.text((mx + x2) / 2, (my + y2) / 2 + 1.5, card.split(":")[-1], fontsize=8, ha="center")

    for name, (x, y) in entities.items():
        rect(x, y, name)

    connect("用户", "拥有", "角色", "M:N", 0, 6)
    connect("角色", "授权", "菜单", "M:N", 0, 4)
    connect("实验室", "包含", "用户", "1:N", -4, 0)
    connect("实验室", "拥有", "设备", "1:N", -3, 3)
    connect("实验室", "拥有", "耗材", "1:N", 3, 0)
    connect("实验室", "拥有", "危化品", "1:N", 6, -4)
    connect("实验室", "管理", "库存", "1:N", 8, -10)
    connect("设备类别", "分类", "设备", "1:N", 0, 4)
    connect("耗材类别", "分类", "耗材", "1:N", 0, 4)
    connect("设备", "借用", "设备借用", "1:N", -5, 0)
    connect("设备", "维修", "设备维修", "1:N", -11, 0)
    connect("设备", "校准", "设备校准", "1:N", -17, 0)
    connect("用户", "申请", "设备借用", "1:N", 10, -8)
    connect("用户", "产生", "系统日志", "1:N", -2, -6)
    connect("耗材", "入库", "耗材入库", "1:N", -2, -4)
    connect("耗材", "出库", "耗材出库", "1:N", 2, -4)
    connect("危化品", "使用", "危化品使用", "1:N", -8, 0)
    connect("耗材", "形成批次", "库存", "1:N", -4, -6)
    connect("危化品", "形成批次", "库存", "1:N", -2, 6)

    ax.text(50, 96, "实验室设备与耗材管理系统总体 Chen E-R 图", ha="center", fontsize=16, fontweight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=240, bbox_inches="tight")
    plt.close(fig)


def draw_entity_attribute_diagrams(output_dir: Path) -> list[Path]:
    plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    paths: list[Path] = []

    for idx, (entity, attrs) in enumerate(ENTITY_ATTRIBUTES.items(), start=1):
        fig, ax = plt.subplots(figsize=(12, 7))
        ax.set_xlim(0, 100)
        ax.set_ylim(0, 100)
        ax.axis("off")

        center = (50, 50)
        ax.add_patch(Rectangle((40, 45), 20, 10, fill=False, linewidth=1.8))
        ax.text(center[0], center[1], entity, ha="center", va="center", fontsize=13, fontweight="bold")

        radius_x = 32
        radius_y = 25
        angles = [2 * math.pi * i / len(attrs) for i in range(len(attrs))]
        for angle, attr in zip(angles, attrs):
            ex = center[0] + radius_x * math.cos(angle)
            ey = center[1] + radius_y * math.sin(angle)
            ax.add_patch(Ellipse((ex, ey), 18, 8, fill=False, linewidth=1.2))
            ax.text(ex, ey, attr, ha="center", va="center", fontsize=10)
            vx = center[0] + 10 * math.cos(angle)
            vy = center[1] + 5 * math.sin(angle)
            ax.plot([vx, ex - 9 * math.cos(angle)], [vy, ey - 4 * math.sin(angle)], color="black", linewidth=1)

        ax.text(50, 92, f"{entity} 属性图", ha="center", fontsize=15, fontweight="bold")
        fig.tight_layout()
        out = output_dir / f"entity_attr_{idx:02d}.png"
        fig.savefig(out, dpi=220, bbox_inches="tight")
        plt.close(fig)
        paths.append(out)
    return paths


def add_relation_table(doc: Document) -> None:
    rows = [["实体A", "联系", "实体B", "基数"]]
    rows.extend([list(item) for item in RELATIONS])
    add_table(doc, rows)


def add_entity_attribute_summary(doc: Document) -> None:
    rows = [["实体", "核心属性"]]
    for entity, attrs in ENTITY_ATTRIBUTES.items():
        rows.append([entity, "、".join(attrs)])
    add_table(doc, rows)


def set_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_begin)
    paragraph._p.append(instr)
    paragraph._p.append(fld_end)


def build_docx() -> Path:
    ensure_assets_dir()
    text = SOURCE_MD.read_text(encoding="utf-8")
    sections = parse_markdown_sections(text)
    section_map = get_section_map(sections)

    er_path = ASSET_DIR / "system_chen_er.png"
    draw_system_er(er_path)
    entity_paths = draw_entity_attribute_diagrams(ASSET_DIR)

    doc = Document()
    set_document_style(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    set_page_number(section)

    add_cover_page(doc)
    add_doc_intro(doc)

    add_heading(doc, "1.1 系统描述", 2)
    add_paragraph(
        doc,
        "实验室设备与耗材管理系统面向高校实验室日常教学、科研与安全管理场景，统一管理设备、耗材、危化品、库存、借还、维修、校准、提醒与审计等核心业务。系统通过数据库驱动的业务台账和权限控制机制，解决实验室资产分散、状态难追踪、库存不透明和安全管理链路不完整等问题。",
    )
    add_paragraph(
        doc,
        "系统采用前后端分离架构，后端基于 Spring Boot 3、Spring Security、JWT 和 MyBatis-Plus 实现业务服务与权限控制，前端基于 Vue 3、TypeScript、Element Plus 和 ECharts 实现管理界面，数据库采用 MySQL 8，缓存采用 Redis 7。",
    )

    for numbered, title in [
        ("1、系统目标", "系统目标"),
        ("2、用户角色分析", "用户角色分析"),
        ("3、功能需求", "功能需求"),
        ("4、非功能需求", "非功能需求"),
        ("5、核心业务流程", "核心业务流程"),
    ]:
        add_paragraph(doc, numbered, first_line_indent=False)
        add_section_from_md(doc, title, get_body_text(section_map.get(title)))

    add_heading(doc, "1.2 用例图", 2)
    add_paragraph(
        doc,
        "本系统的用例图应围绕“实验室资产管理”和“实验室安全与库存管理”两条主线展开，参与者主要包括系统管理员、实验室主任、设备管理员、耗材管理员、危化品管理员、教师、学生、维修人员和校准人员。",
    )
    add_paragraph(
        doc,
        "核心用例包括：用户登录、用户与角色管理、实验室档案维护、设备分类管理、设备台账管理、设备借用与归还、设备维修、设备校准、耗材入库、耗材出库、危化品领用与归还、废液处理、库存预警、报表统计和审计日志查询。绘制正式作业图时，可按“前台申请类用例”和“后台管理类用例”分组展示。",
    )

    add_heading(doc, "1.3 用例说明书", 2)
    use_case_rows = [
        ["项目", "内容"],
        ["用例名称", "设备借用"],
        ["参与者", "教师、学生、设备管理员、实验室主任"],
        ["前置条件", "借用人已登录系统；目标设备状态为“可用”；借用时间和应还时间填写完整。"],
        ["后置条件", "借用申请生成；审批通过后形成借用记录；设备状态更新为“借出”；归还后恢复为“可用”。"],
        ["基本事件流", "1. 借用人选择设备并填写借用用途、借用时间和应还时间。2. 系统校验设备状态及时间合法性。3. 高价值设备提交实验室主任审批。4. 审批通过后登记借出并更新设备状态。5. 归还时管理员核验设备状态并完成归还登记。"],
        ["其他事件流", "1. 借用人主动撤销未审批申请。2. 管理员可在审批前补充借用备注或调整应还日期。"],
        ["异常事件流", "1. 设备当前不可借用时系统拒绝提交。2. 应还时间早于借用时间时提示修正。3. 逾期未归还时系统生成催还提醒。"],
        ["用例关系", "可被“设备台账管理”包含；与“审计日志记录”存在扩展关系。"],
    ]
    add_table(doc, use_case_rows)

    add_heading(doc, "1.4 数据库设计要求", 2)
    requirements = [
        "在现有需求分析基础上，对实验室设备、耗材、危化品、库存、审批和日志等业务对象进行适度抽象与补充，保证概念模型完整。",
        "采用经典 E-R 模型（Chen 模型）绘制系统总体 E-R 图，在总体图中只描述实体与联系，不展开属性。",
        "针对各个核心实体分别绘制属性图，明确主属性、描述属性和关键业务属性。",
        "在数据库逻辑设计中落实主键、外键、非空、唯一、缺省值、检查约束和索引设计，并根据系统功能补充视图、存储过程和触发器。",
        "形成可直接提交的 Word 文档，作为课程设计“需求分析与数据库概念结构设计”阶段成果。",
    ]
    for idx, text_item in enumerate(requirements, start=1):
        add_paragraph(doc, f"（{idx}）{text_item}", first_line_indent=False)

    add_heading(doc, "1.5 系统总体 E-R 图（Chen 模型）", 2)
    add_paragraph(
        doc,
        "下图采用经典 Chen 模型描述系统总体概念结构，仅展示实体、联系及基数，不在总体图中展开属性。矩形表示实体，菱形表示联系，联系两端通过文字标注 1、N 或 M 表示基数。",
    )
    doc.add_picture(str(er_path), width=Cm(16.5))
    pic_p = doc.paragraphs[-1]
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, "图 1-1 实验室设备与耗材管理系统总体 E-R 图", first_line_indent=False)
    add_relation_table(doc)

    add_heading(doc, "1.6 各实体属性图", 2)
    add_paragraph(
        doc,
        "在总体 E-R 图的基础上，对系统核心实体进一步绘制属性图。属性图仍采用 Chen 模型表示，其中矩形为实体，椭圆为属性，便于说明各实体的数据组成。",
    )
    add_entity_attribute_summary(doc)

    for path in entity_paths:
        doc.add_picture(str(path), width=Cm(15.5))
        p = doc.paragraphs[-1]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label = path.stem.replace("entity_attr_", "图 1-").lstrip("0")
        add_paragraph(doc, f"{label} 核心实体属性图", first_line_indent=False)

    add_heading(doc, "1.7 数据库逻辑设计补充说明", 2)
    for title in ["数据库逻辑设计规范", "视图、存储过程与触发器", "数据库设计特点", "当前已实现功能", "推荐后续扩展"]:
        add_paragraph(doc, title, first_line_indent=False)
        add_section_from_md(doc, title, get_body_text(section_map.get(title)))

    if len(doc.sections) == 1:
        doc.add_section(WD_SECTION.NEW_PAGE)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    out = build_docx()
    print(out)
